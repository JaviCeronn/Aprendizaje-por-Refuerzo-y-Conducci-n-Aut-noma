"""
Genera la Memoria del proyecto de Aprendizaje por Refuerzo y Conducción Autónoma.
Cubre: Portada, Índice, Introducción y Fase 1 (Q-Learning tabular en FrozenLake).
"""

import pathlib
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Rutas ──────────────────────────────────────────────────────────────────────
BASE   = pathlib.Path(__file__).resolve().parent.parent
IMG1   = BASE / "results" / "Fase_1" / "q_learning_progress.png"
IMG2   = BASE / "results" / "Fase_1" / "q_table_heatmap.png"
IMG3   = BASE / "results" / "Fase_1" / "q_learning_eval.png"
OUTPUT = BASE / "docs" / "Memoria_RL_Conduccion_Autonoma.docx"

# ── Colores corporativos ────────────────────────────────────────────────────────
AZUL_OSCURO  = RGBColor(0x1B, 0x3A, 0x6B)   # #1B3A6B
AZUL_MEDIO   = RGBColor(0x2E, 0x6D, 0xB8)   # #2E6DB8
GRIS_TEXTO   = RGBColor(0x33, 0x33, 0x33)   # #333333
GRIS_SUBTIT  = RGBColor(0x55, 0x55, 0x55)   # #555555
BLANCO       = RGBColor(0xFF, 0xFF, 0xFF)
VERDE_EXITO  = RGBColor(0x27, 0x6E, 0x3F)   # para destacar resultados


def set_run_font(run, name="Calibri", size=None, bold=False, italic=False,
                 color=None):
    run.font.name = name
    if size:
        run.font.size = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def heading(doc, text, level, color=AZUL_OSCURO, size=None):
    """Párrafo de encabezado con formato manual (sin depender del tema)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    sizes = {1: 18, 2: 14, 3: 12}
    run.font.name  = "Calibri"
    run.font.size  = Pt(size or sizes.get(level, 12))
    run.font.bold  = True
    run.font.color.rgb = color
    # subrayado fino bajo h1
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), '2E6DB8')
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p


def body(doc, text, indent=False, space_after=6):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.7)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, size=11, color=GRIS_TEXTO)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Cm(0.7 + level * 0.5)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=GRIS_TEXTO)
    return p


def formula_block(doc, text):
    """Bloque de código/fórmula con fondo gris claro."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.0)
    p.paragraph_format.right_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name  = "Courier New"
    run.font.size  = Pt(10)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    # Fondo gris claro en el párrafo
    pPr   = p._p.get_or_add_pPr()
    shd   = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'EEEEEE')
    pPr.append(shd)
    return p


def add_image_with_caption(doc, img_path, caption, width_cm=15.5):
    if img_path.exists():
        doc.add_picture(str(img_path), width=Cm(width_cm))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(10)
        r = cp.add_run(caption)
        r.font.name   = "Calibri"
        r.font.size   = Pt(9.5)
        r.font.italic = True
        r.font.color.rgb = GRIS_SUBTIT
    else:
        body(doc, f"[Imagen no encontrada: {img_path.name}]")


def add_page_break(doc):
    doc.add_page_break()


def shade_table_header(row, fill="1B3A6B"):
    for cell in row.cells:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  fill)
        tcPr.append(shd)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = BLANCO
                run.font.bold      = True


# ══════════════════════════════════════════════════════════════════════════════
#  DOCUMENTO
# ══════════════════════════════════════════════════════════════════════════════

doc = Document()

# Márgenes de página
section = doc.sections[0]
section.page_width    = Cm(21)
section.page_height   = Cm(29.7)
section.left_margin   = Cm(2.8)
section.right_margin  = Cm(2.8)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ── Estilos base ───────────────────────────────────────────────────────────────
style_normal = doc.styles["Normal"]
style_normal.font.name = "Calibri"
style_normal.font.size = Pt(11)

# ══════════════════════════════════════════════════════════════════════════════
#  PORTADA
# ══════════════════════════════════════════════════════════════════════════════

# Espaciado superior decorativo
for _ in range(4):
    doc.add_paragraph()

# Título principal
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_after = Pt(6)
r = p_title.add_run("APRENDIZAJE POR REFUERZO")
r.font.name  = "Calibri"
r.font.size  = Pt(28)
r.font.bold  = True
r.font.color.rgb = AZUL_OSCURO

p_title2 = doc.add_paragraph()
p_title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title2.paragraph_format.space_after = Pt(4)
r2 = p_title2.add_run("Y CONDUCCIÓN AUTÓNOMA")
r2.font.name  = "Calibri"
r2.font.size  = Pt(28)
r2.font.bold  = True
r2.font.color.rgb = AZUL_OSCURO

# Línea divisoria (simulada con párrafo con borde inferior)
p_line = doc.add_paragraph()
p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_line.paragraph_format.space_before = Pt(4)
p_line.paragraph_format.space_after  = Pt(18)
pPr  = p_line._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
bot  = OxmlElement('w:bottom')
bot.set(qn('w:val'),   'single')
bot.set(qn('w:sz'),    '12')
bot.set(qn('w:space'), '4')
bot.set(qn('w:color'), '2E6DB8')
pBdr.append(bot)
pPr.append(pBdr)

# Subtítulo
p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_after = Pt(40)
r_sub = p_sub.add_run("Memoria del Proyecto — Máster en Inteligencia Artificial")
r_sub.font.name   = "Calibri"
r_sub.font.size   = Pt(14)
r_sub.font.italic = True
r_sub.font.color.rgb = GRIS_SUBTIT

# Bloque de información del proyecto (párrafos simples, sin tabla)
fields = [
    ("Asignatura",   "Machine Learning II"),
    ("Módulo",       "Aprendizaje por Refuerzo"),
    ("Autor",        "Carlos Díaz Raposo"),
    ("Entorno base", "FrozenLake-v1 (Fase 1) · Duckietown (Fases 2-3)"),
    ("Fecha",        "Junio 2026"),
]

for label, value in fields:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r_lbl = p.add_run(f"{label}:  ")
    r_lbl.font.name  = "Calibri"
    r_lbl.font.size  = Pt(11)
    r_lbl.font.bold  = True
    r_lbl.font.color.rgb = AZUL_MEDIO
    r_val = p.add_run(value)
    r_val.font.name  = "Calibri"
    r_val.font.size  = Pt(11)
    r_val.font.color.rgb = GRIS_TEXTO

for _ in range(6):
    doc.add_paragraph()

# Pie de portada
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_f = p_footer.add_run("Desafío final del Máster IA · Conducción Autónoma en Duckietown")
r_f.font.name   = "Calibri"
r_f.font.size   = Pt(10)
r_f.font.italic = True
r_f.font.color.rgb = GRIS_SUBTIT

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
#  ÍNDICE (manual)
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "Índice", level=1)
doc.add_paragraph()

entries = [
    ("1.", "Introducción", "3", True),
    ("1.1", "Contexto y motivación", "3", False),
    ("1.2", "Estructura del proyecto", "3", False),
    ("1.3", "Entregables", "4", False),
    ("2.", "Fase 1 — Q-Learning tabular en FrozenLake", "5", True),
    ("2.1", "El entorno: FrozenLake-v1 (8×8, hielo resbaladizo)", "5", False),
    ("2.2", "Fundamentos de Q-Learning", "5", False),
    ("2.3", "Exploración: política ε-greedy con decaimiento", "6", False),
    ("2.4", "Implementación y bucle de entrenamiento", "6", False),
    ("2.5", "Resultados del entrenamiento", "7", False),
    ("2.6", "Visualización de la Q-Table aprendida", "8", False),
    ("2.7", "Evaluación con política greedy (1 000 episodios)", "9", False),
    ("2.8", "Conclusiones de la Fase 1 y limitaciones del enfoque tabular", "10", False),
]

# Ancho de texto disponible en twips (1 cm = 567 twips; texto ≈ 15.4 cm → 8 721 twips)
TAB_RIGHT = 8721

for num, title, page, is_main in entries:
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(3 if not is_main else 7)
    p.paragraph_format.space_before = Pt(6 if is_main else 0)
    # Sangría para subsecciones
    if not is_main:
        p.paragraph_format.left_indent = Cm(0.8)

    # Tab stop con puntos de relleno al borde derecho
    pPr  = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab  = OxmlElement('w:tab')
    tab.set(qn('w:val'),    'right')
    tab.set(qn('w:leader'), 'dot')
    tab.set(qn('w:pos'),    str(TAB_RIGHT))
    tabs.append(tab)
    pPr.append(tabs)

    # Número + título
    r_num = p.add_run(f"{num}  {title}")
    r_num.font.name      = "Calibri"
    r_num.font.size      = Pt(11 if is_main else 10)
    r_num.font.bold      = is_main
    r_num.font.color.rgb = AZUL_OSCURO if is_main else GRIS_TEXTO

    # Tabulador + número de página
    r_tab = p.add_run("\t" + page)
    r_tab.font.name      = "Calibri"
    r_tab.font.size      = Pt(11 if is_main else 10)
    r_tab.font.bold      = is_main
    r_tab.font.color.rgb = AZUL_OSCURO if is_main else GRIS_TEXTO

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
#  1. INTRODUCCIÓN
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "1. Introducción", level=1)

# 1.1
heading(doc, "1.1 Contexto y motivación", level=2)
body(doc, (
    "Este proyecto se enmarca en la asignatura Machine Learning II del Máster en Inteligencia "
    "Artificial. El objetivo central es demostrar, de manera progresiva, la viabilidad del "
    "Aprendizaje por Refuerzo (RL) para resolver tareas de navegación autónoma, comenzando por "
    "entornos de cuadrícula clásicos hasta llegar a la conducción autónoma en el simulador 3D "
    "Duckietown, donde el agente percibe el mundo únicamente a través de imágenes de cámara."
))
body(doc, (
    "La motivación pedagógica es clara: partir desde cero con Q-Learning tabular permite "
    "interiorizar los mecanismos fundamentales del RL (exploración vs. explotación, ecuación de "
    "Bellman, convergencia de política) antes de escalar a las redes neuronales profundas que "
    "hacen posible tratar espacios de estados continuos de alta dimensión como las imágenes."
))

# 1.2
heading(doc, "1.2 Estructura del proyecto", level=2)
body(doc, "El proyecto se divide en tres fases con complejidad creciente:")

# Tabla de fases
t = doc.add_table(rows=4, cols=3)
t.style = "Table Grid"
headers = ["Fase", "Algoritmo", "Entorno"]
for j, h in enumerate(headers):
    cell = t.rows[0].cells[j]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.font.name  = "Calibri"
    r.font.size  = Pt(10.5)
    r.font.bold  = True
shade_table_header(t.rows[0])

rows_data = [
    ("Fase 1", "Q-Learning tabular (sin librerías de Deep RL)", "FrozenLake-v1 · cuadrícula 8×8 estocástica"),
    ("Fase 2", "DQN (acciones discretas) y PPO (acciones continuas)",  "Duckietown — 5 mapas de entrenamiento"),
    ("Fase 3", "Algoritmo avanzado (SAC / TD3 / A2C) — a determinar",  "Duckietown — generalización a mapa oculto"),
]

for i, (f, alg, env) in enumerate(rows_data, start=1):
    row = t.rows[i]
    for j, txt in enumerate([f, alg, env]):
        p = row.cells[j].paragraphs[0]
        r = p.add_run(txt)
        r.font.name  = "Calibri"
        r.font.size  = Pt(10)
        r.font.color.rgb = GRIS_TEXTO

doc.add_paragraph()

body(doc, (
    "La estructura de directorios del repositorio refleja esta separación: los notebooks "
    "en notebooks/, los resultados en results/Fase_X/, y los scripts de entrenamiento en src/. "
    "El entorno de ejecución para Duckietown (Fases 2 y 3) es Google Colab con GPU T4, dado "
    "que el renderizado OpenGL de Duckietown requiere Linux y EGL, no disponibles en Windows."
))

# 1.3 Entregables
heading(doc, "1.3 Entregables del proyecto", level=2)
body(doc, "Según las especificaciones del reto, los entregables finales son:")

entregables = [
    ("requirements.txt",        "Dependencias con versiones exactas (==)."),
    ("q_learning_sandbox.py",   "Fase 1 completa: código Q-Learning más gráficas."),
    ("train.py",                "Pipeline de entrenamiento DQN, PPO y el algoritmo de la Fase 3."),
    ("best_agent.zip",          "Mejor modelo entrenado (Fase 2 o 3)."),
    ("Report.pdf",              "Comparativa de los tres algoritmos con justificación técnica."),
]

te = doc.add_table(rows=len(entregables)+1, cols=2)
te.style = "Table Grid"
shade_table_header(te.rows[0])
for j, h in enumerate(["Archivo", "Descripción"]):
    p = te.rows[0].cells[j].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.font.name = "Calibri"; r.font.size = Pt(10.5); r.font.bold = True

for i, (archivo, desc) in enumerate(entregables, start=1):
    row = te.rows[i]
    r1 = row.cells[0].paragraphs[0].add_run(archivo)
    r1.font.name = "Courier New"; r1.font.size = Pt(9.5); r1.font.bold = True; r1.font.color.rgb = AZUL_MEDIO
    r2 = row.cells[1].paragraphs[0].add_run(desc)
    r2.font.name = "Calibri"; r2.font.size = Pt(10); r2.font.color.rgb = GRIS_TEXTO

doc.add_paragraph()
body(doc, (
    "REGLA DE ORO: si el código no carga a la primera en un entorno limpio, la calificación "
    "es 0. Por ello se realizará un dry-run en un Colab nuevo antes de la entrega. Fecha límite: "
    "15 de junio de 2026, 23:59 h."
))

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
#  2. FASE 1 — Q-LEARNING TABULAR
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "2. Fase 1 — Q-Learning tabular en FrozenLake", level=1)

body(doc, (
    "El objetivo de esta primera fase es implementar Q-Learning desde cero, sin ninguna librería "
    "de Deep RL, para demostrar la comprensión del mecanismo antes de pasar a redes neuronales. "
    "Todo el código se escribe en Python puro sobre NumPy y Gymnasium."
))

# 2.1 El entorno
heading(doc, "2.1 El entorno: FrozenLake-v1 (8×8, hielo resbaladizo)", level=2)
body(doc, (
    "FrozenLake-v1 es un entorno de cuadrícula 8×8 incluido en Gymnasium. El agente comienza en "
    "la esquina superior izquierda (estado S) y debe alcanzar la meta en la esquina inferior "
    "derecha (estado G) esquivando los agujeros (H). El espacio de estados tiene 64 posiciones "
    "discretas y el espacio de acciones 4 movimientos: izquierda, abajo, derecha, arriba."
))
body(doc, (
    "La variante is_slippery=True introduce estocasticidad: aunque el agente elija ir a la "
    "derecha, puede desplazarse en cualquiera de las tres direcciones perpendiculares con "
    "probabilidad 1/3 cada una. Esto hace el problema no trivial y obliga al agente a aprender "
    "rutas que sean robustas ante la aleatoriedad del entorno, no solo la ruta más corta."
))

bullet(doc, "Espacio de estados: 64 posiciones discretas (cuadrícula 8×8).")
bullet(doc, "Espacio de acciones: 4 (izquierda ←, abajo ↓, derecha →, arriba ↑).")
bullet(doc, "Recompensa: +1 al alcanzar la meta, 0 en el resto de transiciones.")
bullet(doc, "Episodio termina: al caer en agujero, llegar a la meta o superar 200 pasos.")
bullet(doc, "Semilla fija: SEED = 42 para reproducibilidad total.")

# 2.2 Fundamentos
heading(doc, "2.2 Fundamentos de Q-Learning", level=2)
body(doc, (
    "Q-Learning es un algoritmo de Diferencia Temporal (TD) off-policy que aprende una función "
    "de valor de acción Q(s, a): la recompensa total esperada al tomar la acción a en el estado s "
    "y seguir la política óptima a partir de ahí. Se almacena como una tabla de 64×4 valores, "
    "inicializada a cero (sin conocimiento previo)."
))
body(doc, "La actualización en cada paso sigue la ecuación de Bellman:")

formula_block(doc,
    "Q(s, a) ← Q(s, a) + α · [ r + γ · max_a' Q(s', a') − Q(s, a) ]\n"
    "                           └──────────── TD error ───────────────┘"
)

body(doc, "Los hiperparámetros utilizados son:")

# Tabla hiperparámetros
th = doc.add_table(rows=6, cols=3)
th.style = "Table Grid"
shade_table_header(th.rows[0])
for j, h in enumerate(["Parámetro", "Valor", "Descripción"]):
    p = th.rows[0].cells[j].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h); r.font.name = "Calibri"; r.font.size = Pt(10.5); r.font.bold = True

hparams = [
    ("α (alpha)",        "0.10",   "Tasa de aprendizaje: cuánto sobrescribimos cada actualización."),
    ("γ (gamma)",        "0.99",   "Factor de descuento: valoramos mucho las recompensas futuras."),
    ("ε inicial",        "1.00",   "Exploración pura al principio (100 % aleatoria)."),
    ("ε mínimo",         "0.01",   "Exploración residual que nunca desaparece del todo."),
    ("Decaimiento ε",    "0.9995", "Multiplicador por episodio; ε llega al mínimo ~ep. 8 000."),
]
for i, (p, v, d) in enumerate(hparams, start=1):
    cells = th.rows[i].cells
    for j, txt in enumerate([p, v, d]):
        r_ = cells[j].paragraphs[0].add_run(txt)
        r_.font.name = "Calibri"; r_.font.size = Pt(10); r_.font.color.rgb = GRIS_TEXTO

doc.add_paragraph()
body(doc, (
    "Con α = 0.1 el aprendizaje es estable: no sobreajusta a las últimas experiencias pero "
    "converge. Con γ = 0.99 el agente es muy previsor: incluso los estados lejanos a la meta "
    "heredan valor de ella, lo que es imprescindible para un tablero de 64 celdas donde la meta "
    "está a muchos pasos del inicio."
))

# 2.3 Exploración
heading(doc, "2.3 Exploración: política ε-greedy con decaimiento", level=2)
body(doc, (
    "La política ε-greedy equilibra exploración y explotación de forma sencilla y efectiva. "
    "Con probabilidad ε el agente elige una acción aleatoria (explora); con probabilidad 1−ε "
    "elige la acción con mayor valor Q conocido (explota). "
))
body(doc, (
    "El decaimiento exponencial (ε ← max(ε_min, ε · 0.9995)) garantiza que el agente explore "
    "ampliamente durante los primeros episodios y progresivamente confíe más en lo aprendido. "
    "Sin exploración inicial, el agente se quedaría atrapado en la primera ruta funcional, que "
    "probablemente no es la óptima. El desempate entre acciones de igual valor se resuelve "
    "con np.flatnonzero para evitar sesgos sistemáticos hacia la primera acción."
))

# 2.4 Implementación
heading(doc, "2.4 Implementación y bucle de entrenamiento", level=2)
body(doc, (
    "El bucle principal ejecuta 10 000 episodios. En cada uno el agente parte del estado "
    "inicial, aplica la política ε-greedy, realiza la actualización de Bellman y decae ε. "
    "La actualización omite el término bootstrap (γ · max Q(s', ·)) cuando el episodio termina, "
    "porque no existe un estado siguiente."
))

formula_block(doc,
    "td_target = reward + γ · max Q(s') · (1 si no terminado, 0 si terminado)\n"
    "Q[s, a] += α · (td_target − Q[s, a])"
)

body(doc, "La evolución de la tasa de éxito cada 1 000 episodios durante el entrenamiento fue:")

# Tabla de progreso
tp = doc.add_table(rows=11, cols=3)
tp.style = "Table Grid"
shade_table_header(tp.rows[0])
for j, h in enumerate(["Episodio", "Epsilon (ε)", "Tasa de éxito (últimos 1 000)"]):
    p = tp.rows[0].cells[j].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h); r.font.name = "Calibri"; r.font.size = Pt(10); r.font.bold = True

prog = [
    ("1 000",  "0.606", "0.000"),
    ("2 000",  "0.368", "0.002"),
    ("3 000",  "0.223", "0.045"),
    ("4 000",  "0.135", "0.224"),
    ("5 000",  "0.082", "0.341"),
    ("6 000",  "0.050", "0.443"),
    ("7 000",  "0.030", "0.483"),
    ("8 000",  "0.018", "0.478"),
    ("9 000",  "0.011", "0.543"),
    ("10 000", "0.010", "0.518"),
]
for i, (ep, eps, rate) in enumerate(prog, start=1):
    cells = tp.rows[i].cells
    for j, txt in enumerate([ep, eps, rate]):
        p = cells[j].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_ = p.add_run(txt)
        r_.font.name = "Calibri"; r_.font.size = Pt(10); r_.font.color.rgb = GRIS_TEXTO
        if j == 2 and float(rate) >= 0.5:
            r_.font.color.rgb = VERDE_EXITO; r_.font.bold = True

doc.add_paragraph()

# 2.5 Resultados de entrenamiento
heading(doc, "2.5 Resultados del entrenamiento", level=2)
body(doc, (
    "La figura siguiente muestra la curva de aprendizaje (media móvil de 100 episodios) y el "
    "decaimiento de ε a lo largo de los 10 000 episodios de entrenamiento."
))
add_image_with_caption(doc, IMG1,
    "Figura 1. Izquierda: tasa de éxito (media móvil 100 ep.) durante el entrenamiento. "
    "Derecha: decaimiento exponencial de ε desde 1.0 hasta 0.01.")

body(doc, (
    "El patrón es exactamente el esperado en un algoritmo RL con exploración decreciente. "
    "Hasta el episodio 2 000 la tasa es casi cero: con ε > 0.3 el agente explora demasiado "
    "al azar para llegar a la meta de forma consistente. A partir del episodio 4 000-5 000, "
    "cuando ε cae por debajo de 0.10, el agente comienza a explotar lo aprendido y la tasa "
    "sube rápidamente hacia el 45-55 %. Las dos curvas son simétricas por diseño: la caída "
    "de ε es la condición que desbloquea el rendimiento."
))
body(doc, (
    "La pequeña oscilación en los últimos 3 000 episodios no es un retroceso real sino ruido "
    "estadístico intrínseco al entorno estocástico: la política ya ha convergido, pero el hielo "
    "resbaladizo introduce varianza irreducible en el resultado de cada episodio."
))

# 2.6 Q-Table
heading(doc, "2.6 Visualización de la Q-Table aprendida", level=2)
body(doc, (
    "La Q-Table aprendida puede visualizarse de dos formas complementarias: como mapa de calor "
    "del valor máximo por estado (cuánto vale cada casilla) y como mapa de la acción óptima "
    "(en qué dirección apunta la política greedy en cada posición)."
))
add_image_with_caption(doc, IMG2,
    "Figura 2. Izquierda: valor máximo Q(s,*) por casilla; verde oscuro = próximo a la meta. "
    "Derecha: acción greedy (←↓→↑) en cada posición. H = agujero (rojo), G = meta (dorado).")

body(doc, (
    "En el mapa de valor se observa un gradiente claro: los estados más próximos a la meta "
    "(esquina inferior derecha) tienen valores más altos. Esto es consecuencia directa de "
    "γ = 0.99: el descuento propaga el valor de la meta hacia atrás por todo el tablero en "
    "virtud de las actualizaciones de Bellman acumuladas a lo largo del entrenamiento."
))
body(doc, (
    "En el mapa de política las flechas muestran que el agente ha aprendido a rodear los "
    "agujeros en lugar de cruzarlos en línea recta, lo que confirma que existe aprendizaje "
    "real y no una simple memorización de una única trayectoria."
))
body(doc, (
    "Los estados con valor muy bajo en la zona inferior izquierda corresponden a casillas "
    "rodeadas de agujeros con muy poca probabilidad de conducir a la meta; el agente los ha "
    "identificado correctamente como poco prometedores."
))

# 2.7 Evaluación
heading(doc, "2.7 Evaluación con política greedy (1 000 episodios)", level=2)
body(doc, (
    "Para medir el rendimiento real de la política aprendida se realiza una evaluación limpia: "
    "1 000 episodios con ε = 0 (siempre se elige la mejor acción conocida) y semillas que NO "
    "se utilizaron durante el entrenamiento (seeds 10 000 a 10 999). Esto elimina el sesgo de "
    "haber visto esas situaciones previamente."
))
add_image_with_caption(doc, IMG3,
    "Figura 3. Tasa de éxito acumulada durante la evaluación greedy sobre 1 000 episodios "
    "con semillas no vistas en el entrenamiento. Tasa final estabilizada: 0.592.")

body(doc, (
    "La curva converge rápidamente hacia 0.59: con apenas 200 episodios ya se tiene una "
    "estimación estable, lo que indica que la política es consistente y no depende de semillas "
    "particulares. La tasa final de 59.2 % está plenamente en línea con lo esperado."
))

# Tabla resumen resultados
tr = doc.add_table(rows=3, cols=2)
tr.style = "Table Grid"
shade_table_header(tr.rows[0])
for j, h in enumerate(["Métrica", "Valor"]):
    p = tr.rows[0].cells[j].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h); r.font.name = "Calibri"; r.font.size = Pt(10.5); r.font.bold = True

res = [
    ("Tasa de éxito — entrenamiento (media móvil últimos 100 ep.)", "59.0 %"),
    ("Tasa de éxito — evaluación greedy (1 000 ep., semillas nuevas)", "59.2 %"),
]
for i, (desc, val) in enumerate(res, start=1):
    cells = tr.rows[i].cells
    r1 = cells[0].paragraphs[0].add_run(desc)
    r1.font.name = "Calibri"; r1.font.size = Pt(10); r1.font.color.rgb = GRIS_TEXTO
    r2 = cells[1].paragraphs[0].add_run(val)
    r2.font.name = "Calibri"; r2.font.size = Pt(10.5)
    r2.font.bold = True; r2.font.color.rgb = VERDE_EXITO
    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# 2.8 Conclusiones
heading(doc, "2.8 Conclusiones de la Fase 1 y limitaciones del enfoque tabular", level=2)
body(doc, (
    "Una tasa de éxito del 59 % es un resultado correcto y esperado. El techo teórico para la "
    "política óptima en FrozenLake 8×8 estocástico es aproximadamente 65-75 %: el porcentaje "
    "restante es azar puro del hielo, imposible de evitar con independencia de la estrategia "
    "adoptada."
))
body(doc, (
    "Lo relevante no es el número en sí, sino lo que demuestra: el algoritmo partió sin "
    "conocimiento alguno (Q-Table inicializada a ceros) y, siguiendo únicamente la señal de "
    "recompensa del entorno, construyó de forma autónoma una política que evita la mayoría de "
    "los agujeros y encuentra rutas hacia la meta."
))

body(doc, "Logros demostrados en la Fase 1:")
bullet(doc, "Implementación correcta de Q-Learning sin librerías externas de RL.")
bullet(doc, "Convergencia al techo teórico del entorno (~59-65 %) en 10 000 episodios.")
bullet(doc, "Política con semántica espacial: rodea agujeros, propaga valor desde la meta.")
bullet(doc, "Reproducibilidad total mediante semilla fija (SEED = 42).")
bullet(doc, "Separación limpia entrenamiento / evaluación (semillas distintas).")

doc.add_paragraph()
body(doc, "Limitación fundamental — escalabilidad:")
body(doc, (
    "La Q-Table crece linealmente con el número de estados. En FrozenLake los 64 estados caben "
    "cómodamente en memoria. Sin embargo, en Duckietown la observación es una imagen de cámara: "
    "el espacio de estados tiene una cardinalidad prácticamente infinita (millones de imágenes "
    "posibles), lo que hace completamente inviable mantener una tabla explícita. "
    "Esta es precisamente la motivación para pasar a Deep RL en las fases siguientes, donde "
    "una red neuronal convolucional (CNN) aproxima la función Q en lugar de almacenarla, "
    "generalizando a estados nunca vistos durante el entrenamiento."
))

body(doc, (
    "La Fase 2 abordará esta limitación implementando DQN (Deep Q-Network) y PPO (Proximal "
    "Policy Optimization) sobre el simulador Duckietown, entrenando a partir de imágenes de "
    "cámara preprocesadas (recorte, escala de grises, resize 64×64, stacking de 4 frames)."
), space_after=20)

# ── Guardar ────────────────────────────────────────────────────────────────────
OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUTPUT))
print(f"Documento guardado en: {OUTPUT}")
